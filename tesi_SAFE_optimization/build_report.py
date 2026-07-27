"""
Build progress_report_steps10c_10h.docx
Rewrites Steps 10c + 10h in Anila Vata's voice with Unipavia formatting.
"""

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

BASE  = os.path.dirname(os.path.abspath(__file__))
FIG   = os.path.join(BASE, "data", "results", "step10", "frontier_figures")
OUT   = os.path.join(BASE, "data", "results", "progress_report_steps10c_10h.docx")

# colour palette
NAVY  = RGBColor(0x1F, 0x38, 0x64)
BLUE  = RGBColor(0x2E, 0x54, 0x96)
DGREY = RGBColor(0x40, 0x40, 0x40)
BLACK = RGBColor(0x00, 0x00, 0x00)
RED   = RGBColor(0xFF, 0x00, 0x00)

doc = Document()
sec = doc.sections[0]
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, attr, Cm(2.5))

# ── low-level font setter ────────────────────────────────────────────────────
def _f(run, bold=False, sz=12, col=BLACK, italic=False):
    run.font.name      = "Times New Roman"
    run.font.size      = Pt(sz)
    run.font.bold      = bold
    run.font.color.rgb = col
    run.font.italic    = italic

# ── paragraph builders ───────────────────────────────────────────────────────
def body(text, col=BLACK):
    """Body paragraph: 12pt, justified, 2.0 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(24)
    _f(p.add_run(text), col=col)
    return p

def numbered_item(n, text):
    """
    Numbered list item as a body paragraph that starts with the number.
    Keeps the justified body style so it flows with the document.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.left_indent  = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    _f(p.add_run(f"{n}. {text}"))
    return p

def blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(24)
    _f(p.add_run(""))

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(24)
    _f(p.add_run(text), bold=True, sz=14, col=NAVY)

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(24)
    _f(p.add_run(text), bold=True, sz=12, col=BLUE)

def h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(24)
    _f(p.add_run(text), bold=True, sz=12, col=DGREY)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.line_spacing = Pt(20)
    _f(p.add_run(text), sz=10, italic=True)

def embed_fig(fname, cap, width=5.5):
    path = os.path.join(FIG, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width))
    else:
        _f(p.add_run(f"[{cap.split('.')[0]} - file not found, insert manually]"), col=RED)
    caption(cap)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style     = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _f(p.add_run(h), bold=True, sz=10)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(t.rows[ri + 1].cells):
            val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = val.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _f(p.add_run(str(row_data[ci])), sz=10)
    if widths:
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(widths[ci])
    return t

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
blank(); blank(); blank()

for txt, sz, bold in [
    ("Progress Report", 16, True),
    ("SAFE AI Evaluation and Portfolio Optimization", 14, True),
    ("Steps 10c and 10h", 14, True),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _f(p.add_run(txt), bold=bold, sz=sz, col=NAVY)

blank(); blank(); blank()

for line in ["Anila Vata", "MSc Quantitative Finance",
             "University of Pavia", "Supervisor: Prof. Paolo Giudici", "June 2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _f(p.add_run(line))

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 -- STEP 10c
# ─────────────────────────────────────────────────────────────────────────────
h1("Section 1 - Step 10c: SAFE–performance frontier (50×3 configurations)")

# 1.1 ─────────────────────────────────────────────────────────────────────────
h2("1.1 What this step does and why")

body(
    "Step 10c runs the full portfolio optimisation pipeline across 150 model "
    "configurations -- 50 per model family, and records both the SAFE compliance "
    "score and the portfolio performance metrics for each one. Three configurations "
    "per family were enough to verify that the frontier exists, but not enough to "
    "test statistically whether compliance predicts performance or to find "
    "Pareto-dominant points."
)

# 1.2 ─────────────────────────────────────────────────────────────────────────
h2("1.2 Grid construction")

body(
    "The Ridge grid varies the regularisation parameter α across 50 values on a "
    "log scale from 10⁻⁴ to 10⁴. Log spacing is used because the "
    "coefficient structure changes slowly at very small α and changes rapidly near "
    "the transition zone, so equal spacing on the log scale gives more resolution "
    "where it matters."
)
blank()
body("The XGBoost grid draws 50 configurations by random search with seed 42. "
     "The parameters and their ranges are:")
blank()
numbered_item(1, "max_depth ∈ {2, 3, 4, 5, 6, 8}")
numbered_item(2, "learning_rate ∈ {0.005, 0.01, 0.03, 0.05, 0.1, 0.15}")
numbered_item(3, "n_estimators ∈ {50, 100, 150, 200, 300, 500}")
numbered_item(4, "subsample ∈ {0.6, 0.8, 1.0}")
numbered_item(5, "colsample_bytree ∈ {0.6, 0.8, 1.0}")
numbered_item(6, "reg_alpha ∈ {0.0, 0.01, 0.1, 1.0}")
numbered_item(7, "reg_lambda ∈ {0.1, 1.0, 5.0, 10.0}")
blank()
body("The MLP grid also draws 50 configurations by random search with seed 42. "
     "The parameters sampled are:")
blank()
numbered_item(1, "Hidden layer architecture (1–3 layers, widths from 32 to 256 neurons)")
numbered_item(2, "L2 penalty α ∈ {0.0001, 0.001, 0.01, 0.1}")
numbered_item(3, "Initial learning rate ∈ {0.0001, 0.0005, 0.001, 0.005}")
numbered_item(4, "Batch size ∈ {64, 128, 256}")
blank()
body(
    "Activation function and solver are fixed for MLP. All 150 configurations "
    "share the same portfolio settings: SAFE weight λ = 1.0, top-K = 10, "
    "weight bounds [1%, 20%], sector cap 30%, Gurobi solver."
)

# 1.3 ─────────────────────────────────────────────────────────────────────────
h2("1.3 Compliance score calculation")

body(
    "The compliance score (CS) aggregates the four SAFE dimensions - RGA, RGR, "
    "RGE*, RGF* -- into one number. Three aggregation methods are used in parallel "
    "because the choice between them is a normative one, and I want conclusions "
    "that hold across all three."
)
blank()
numbered_item(
    1,
    "Arithmetic mean: the straight average of the four dimension scores. "
    "A weakness on one dimension can be offset by strength on another."
)
numbered_item(
    2,
    "Geometric mean: the fourth root of the product of the four scores. "
    "A single dimension score near zero pulls the whole CS close to zero, "
    "so no dimension can be ignored."
)
numbered_item(
    3,
    "Root-mean-square: penalises configurations whose four scores are uneven, "
    "even if their arithmetic mean is high."
)

# 1.4 ─────────────────────────────────────────────────────────────────────────
h2("1.4 xgboost_47: the Pareto-dominant configuration")

body(
    "Configuration xgboost_47 (max_depth = 5, learning_rate = 0.15, "
    "n_estimators = 500, subsample = 0.6, colsample_bytree = 0.6) achieves "
    "CS = 0.629 (arithmetic), Sharpe = 1.812, MaxDD = 16.3%, and Sortino = 2.599. "
    "No other configuration in the 150-point grid has a higher compliance score "
    "without also having a lower Sharpe. It is therefore strictly Pareto-dominant "
    "in the CS–Sharpe plane."
)
blank()
body(
    "Its one weakness is turnover. Average monthly rebalancing is 92.0%, "
    "the highest in the sample. In a high-transaction-cost setting this must "
    "be factored into any net-of-cost performance comparison."
)
blank()

add_table(
    ["Config ID", "Family", "Key parameters", "CS arith", "Sharpe", "Max DD", "Avg TO", "Sortino"],
    [
        ["xgboost_47", "XGBoost",
         "md=5, lr=0.15, ne=500,\nss=0.6, cb=0.6",
         "0.629", "1.812", "16.3%", "92.0%", "2.599"],
        ["xgboost_43", "XGBoost",
         "md=5, lr=0.05, ne=500,\nss=0.8, cb=0.6",
         "0.613", "1.601", "18.4%", "86.2%", "2.341"],
        ["xgboost_14", "XGBoost",
         "md=6, lr=0.15, ne=150,\nss=1.0, cb=0.6",
         "0.618", "1.321", "18.0%", "89.6%", "2.033"],
        ["xgboost_46", "XGBoost",
         "md=5, lr=0.1, ne=300,\nss=1.0, cb=1.0",
         "0.617", "0.809", "18.1%", "90.6%", "1.212"],
        ["ridge_01",   "Ridge",
         "α = 0.0001",
         "0.549", "0.635", "22.5%", "79.9%", "0.821"],
        ["mlp_15",     "MLP",
         "layers=(64,32),\nα=0.0001, lr=0.001",
         "0.607", "0.498", "17.4%", "96.7%", "0.701"],
        ["mlp_06",     "MLP",
         "layers=(64,32),\nα=0.001, lr=0.0001",
         "0.607", "0.496", "23.2%", "95.6%", "0.723"],
    ],
    widths=[2.4, 2.0, 3.8, 1.8, 1.6, 1.6, 1.6, 1.8],
)
blank()
caption(
    "Table 1. Selected configurations from the 150-point frontier. "
    "xgboost_47 is Pareto-dominant on CS and Sharpe. "
    "md = max_depth, lr = learning_rate, ne = n_estimators, "
    "ss = subsample, cb = colsample_bytree. "
    "Max DD and Avg TO expressed as percentages."
)
blank()

embed_fig(
    "pareto_compliance_score_arithmetic__sharpe.png",
    "Figure 1. Pareto frontier in the compliance score (arithmetic)–Sharpe plane "
    "across 150 configurations. xgboost_47 is the Pareto-dominant point.",
)
blank()
embed_fig(
    "raw_compliance_score_arithmetic__sharpe.png",
    "Figure 2. Scatter of compliance score (arithmetic) vs Sharpe ratio "
    "for all 150 configurations. Colour indicates model family.",
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 - STEP 10h
# ─────────────────────────────────────────────────────────────────────────────
h1("Section 2 -- Step 10h: SAFE–performance scatter analysis")

# 2.1 ─────────────────────────────────────────────────────────────────────────
h2("2.1 What the analysis does")

body(
    "Step 10h tests whether the compliance score is statistically associated "
    "with portfolio performance across the 150-configuration frontier. For each "
    "combination of compliance type (arithmetic, geometric, RMS) and performance "
    "metric (Sharpe, maximum drawdown, average turnover), the analysis computes "
    "a Spearman rank correlation, both pooled across all 150 configurations "
    "and separately within each model family. Each scatter plot adds a LOWESS "
    "curve to show the conditional mean trend without assuming linearity."
)

# 2.2 ─────────────────────────────────────────────────────────────────────────
h2("2.2 Sharpe ratio")

h3("2.2.1 Pooled result")

body(
    "Pooled across all 150 configurations, higher compliance is associated with "
    "higher Sharpe: ρ = 0.597 (arithmetic, p < 0.001), ρ = 0.618 "
    "(geometric, p < 0.001), ρ = 0.417 (RMS, p < 0.001). The relationship "
    "is moderate to strong and holds across all three aggregation methods. "
    "This is the main headline result."
)

h3("2.2.2 Within-family results")

body(
    "The pooled correlation is driven by the gap between model families, not by "
    "variation within them. XGBoost configurations dominate Ridge and MLP on "
    "both CS and Sharpe simultaneously, which inflates the pooled number. "
    "Within XGBoost, ρ = 0.090 (p = 0.537, not significant). Within MLP, "
    "ρ = 0.152 (p = 0.291, not significant). Within Ridge, ρ = 0.498 "
    "(p < 0.001, significant), but this reflects the regularisation path rather "
    "than a clean compliance–return relationship. For a practitioner choosing "
    "among configurations of the same model family, compliance score alone does "
    "not reliably predict which will deliver a higher Sharpe."
)

# 2.3 ─────────────────────────────────────────────────────────────────────────
h2("2.3 Maximum drawdown")

h3("2.3.1 Pooled result")

body(
    "Pooled, the compliance–drawdown correlation is negative and significant: "
    "ρ = −0.350 (arithmetic, p < 0.001). A negative value is the "
    "right direction: higher compliance goes with lower drawdown. The magnitude "
    "is weaker than the Sharpe result."
)

h3("2.3.2 The Ridge anomaly")

body(
    "Within Ridge the correlation is positive: ρ = +0.393 (p = 0.005). "
    "This is an artefact of the regularisation path. At high α, Ridge "
    "shrinks all coefficients towards zero, making the ranking nearly uniform "
    "across stocks. A uniform ranking scores well on robustness and fairness "
    "because there is no cross-group variation to flag, but it also loses "
    "predictive power, so drawdown increases. This is not a genuine alignment "
    "between compliance and risk control."
)
blank()
body(
    "Within MLP, ρ = −0.543 (p < 0.001) -- the strongest within-family "
    "result in the analysis. More compliant MLP configurations genuinely achieve "
    "lower drawdowns. Within XGBoost, ρ = −0.248 (p = 0.082), negative "
    "in direction but not significant."
)

# 2.4 ─────────────────────────────────────────────────────────────────────────
h2("2.4 Average turnover")

h3("2.4.1 Pooled result")

body(
    "Pooled, compliance and turnover are strongly and positively correlated: "
    "ρ = 0.611 (arithmetic, p < 0.001), ρ = 0.598 (geometric), "
    "ρ = 0.669 (RMS). Higher compliance comes with more frequent rebalancing. "
    "This is the main practical trade-off: better fairness and robustness "
    "properties cost more in transaction frequency."
)

h3("2.4.2 Within-family results")

body(
    "Within XGBoost, the trade-off is the sharpest in the sample: "
    "ρ = 0.779 (arithmetic, p < 0.001), ρ = 0.836 (geometric). "
    "The configurations that score highest on compliance within XGBoost are "
    "the most aggressive rebalancers. Within MLP the same direction holds: "
    "ρ = 0.772 (arithmetic, p < 0.001). Within Ridge, the sign reverses "
    "(ρ = −0.339, p = 0.016) because heavy regularisation simultaneously "
    "reduces turnover and inflates some SAFE scores mechanically."
)

# 2.5 ─────────────────────────────────────────────────────────────────────────
h2("2.5 Summary of results")

body("Table 2 summarises the Spearman rank correlations for arithmetic CS.")
blank()

add_table(
    ["Performance metric", "Pooled ρ", "Sig.",
     "Ridge ρ", "Sig.", "XGBoost ρ", "Sig.", "MLP ρ", "Sig."],
    [
        ["Sharpe", "+0.597", "***", "+0.498", "***", "+0.090", "n.s.", "+0.152", "n.s."],
        ["Max DD", "−0.350", "***", "+0.393", "**",
         "−0.248", "n.s.", "−0.543", "***"],
        ["Avg TO", "+0.611", "***", "−0.339", "*",
         "+0.779", "***", "+0.772", "***"],
    ],
    widths=[3.2, 2.0, 1.3, 1.9, 1.3, 2.4, 1.3, 1.9, 1.3],
)
blank()
caption(
    "Table 2. Spearman rank correlations between compliance score (arithmetic) "
    "and performance metrics, pooled and by model family (n = 150 pooled, "
    "n = 50 per family). "
    "*** p < 0.001, ** p < 0.01, * p < 0.05, n.s. not significant. "
    "For Max DD, a negative ρ means higher compliance → lower drawdown."
)
blank()

embed_fig(
    "model_level_safe_sharpe_by_family.png",
    "Figure 3. Compliance score (arithmetic) vs Sharpe ratio by model family. "
    "LOWESS curves fitted separately per family.",
)
blank()
embed_fig(
    "raw_compliance_score_arithmetic__max_drawdown.png",
    "Figure 4. Compliance score (arithmetic) vs maximum drawdown "
    "for all 150 configurations.",
)
blank()
embed_fig(
    "raw_compliance_score_arithmetic__avg_turnover.png",
    "Figure 5. Compliance score (arithmetic) vs average turnover "
    "for all 150 configurations.",
)

# 2.6 ─────────────────────────────────────────────────────────────────────────
h2("2.6 What the results mean")

body("Three findings summarise the analysis:")
blank()

numbered_item(
    1,
    "The pooled compliance–Sharpe correlation (ρ = 0.597) shows that "
    "SAFE-compliant models tend to produce better-performing portfolios when "
    "comparing across model families. Designing a model to be fairer and more "
    "robust does not come at a systematic cost in risk-adjusted return."
)
blank()
numbered_item(
    2,
    "Within a model family, compliance score is not a reliable predictor of "
    "Sharpe ratio. The within-XGBoost and within-MLP correlations are not "
    "statistically significant. A practitioner cannot pick a better-performing "
    "configuration simply by choosing the one with the highest SAFE score "
    "within a given family."
)
blank()
numbered_item(
    3,
    "The compliance–turnover trade-off is real and consistent. "
    "Higher compliance goes with higher rebalancing frequency for both "
    "XGBoost (ρ = 0.779) and MLP (ρ = 0.772). Whether this is "
    "acceptable depends on the transaction cost environment."
)
blank()
body(
    "xgboost_47 is the recommended configuration for deployment. It is "
    "Pareto-dominant on compliance and gross Sharpe, achieves the lowest "
    "drawdown within the XGBoost family, and its compliance advantage is "
    "consistent across all three aggregation methods "
    "(CS arith = 0.629, CS geom = 0.528, CS RMS = 0.661). "
    "Its high turnover should be accounted for in any cost-adjusted evaluation."
)

# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved  : {OUT}")

# validate
from docx import Document as _D

d2   = _D(OUT)
imgs = [p for p in d2.part.package.iter_parts() if "image" in p.content_type]
print(f"Images : {len(imgs)}")
print(f"Tables : {len(d2.tables)}")
print(f"Paras  : {len(d2.paragraphs)}")
size = os.path.getsize(OUT)
print(f"Size   : {size:,} bytes ({size/1024:.1f} KB)")
print("OK")
